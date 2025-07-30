#!/usr/bin/env python3
"""
Extract and analyze AWS and cloud security content from FedRAMP documents
"""

import os
import re
from pathlib import Path
from docx import Document
import openpyxl
import json

class FedRAMPDocumentAnalyzer:
    def __init__(self, base_path):
        self.base_path = Path(base_path)
        self.aws_keywords = [
            'AWS', 'Amazon Web Services', 'EC2', 'S3', 'CloudWatch', 'IAM', 
            'VPC', 'Lambda', 'RDS', 'CloudTrail', 'CloudFormation', 'EBS',
            'ELB', 'Auto Scaling', 'Route 53', 'CloudFront', 'SQS', 'SNS',
            'KMS', 'Secrets Manager', 'Systems Manager', 'GuardDuty', 'Shield',
            'WAF', 'Config', 'Inspector', 'Security Hub', 'SSM', 'ECS', 'EKS',
            'Fargate', 'API Gateway', 'Cognito', 'DynamoDB', 'Redshift'
        ]
        self.cloud_security_keywords = [
            'cloud security', 'cloud infrastructure', 'cloud monitoring',
            'cloud access', 'cloud configuration', 'cloud compliance',
            'cloud encryption', 'cloud backup', 'cloud disaster recovery',
            'cloud authentication', 'cloud authorization', 'cloud logging',
            'cloud alerting', 'cloud vulnerability', 'cloud patch',
            'multi-factor authentication', 'MFA', 'encryption at rest',
            'encryption in transit', 'security group', 'network ACL',
            'bastion host', 'jump box', 'firewall', 'IDS', 'IPS',
            'SIEM', 'SOC', 'incident response', 'access control',
            'privilege', 'least privilege', 'monitoring', 'alerting',
            'audit', 'compliance', 'configuration management'
        ]
        self.findings = {}

    def extract_docx_content(self, file_path):
        """Extract text content from a .docx file"""
        try:
            doc = Document(file_path)
            content = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))
            
            return "\n".join(content)
        except Exception as e:
            return f"Error reading {file_path}: {str(e)}"

    def extract_xlsx_content(self, file_path):
        """Extract text content from a .xlsx file"""
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            content = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                content.append(f"\n=== Sheet: {sheet_name} ===")
                
                for row in sheet.iter_rows():
                    row_data = []
                    for cell in row:
                        if cell.value is not None:
                            row_data.append(str(cell.value))
                    if row_data:
                        content.append(" | ".join(row_data))
            
            return "\n".join(content)
        except Exception as e:
            return f"Error reading {file_path}: {str(e)}"

    def search_content(self, content, file_name):
        """Search content for AWS and cloud security related information"""
        findings = {
            'aws_services': set(),
            'cloud_security_measures': set(),
            'relevant_excerpts': []
        }
        
        # Search for AWS services
        for keyword in self.aws_keywords:
            pattern = rf'\b{re.escape(keyword)}\b'
            if re.search(pattern, content, re.IGNORECASE):
                findings['aws_services'].add(keyword)
        
        # Search for cloud security keywords
        for keyword in self.cloud_security_keywords:
            if keyword.lower() in content.lower():
                findings['cloud_security_measures'].add(keyword)
        
        # Extract relevant excerpts (sentences containing AWS or important security terms)
        sentences = re.split(r'[.!?]\s+', content)
        for sentence in sentences:
            if any(keyword.lower() in sentence.lower() for keyword in self.aws_keywords + ['cloud', 'security', 'monitoring', 'access', 'encryption']):
                # Clean up the sentence
                clean_sentence = ' '.join(sentence.split())
                if len(clean_sentence) > 20 and len(clean_sentence) < 500:
                    findings['relevant_excerpts'].append(clean_sentence)
        
        # Convert sets to lists for JSON serialization
        findings['aws_services'] = list(findings['aws_services'])
        findings['cloud_security_measures'] = list(findings['cloud_security_measures'])
        
        # Limit excerpts to most relevant
        findings['relevant_excerpts'] = findings['relevant_excerpts'][:20]
        
        return findings

    def analyze_document(self, file_path):
        """Analyze a single document"""
        file_path = Path(file_path)
        print(f"\nAnalyzing: {file_path.name}")
        
        if file_path.suffix == '.docx':
            content = self.extract_docx_content(file_path)
        elif file_path.suffix == '.xlsx':
            content = self.extract_xlsx_content(file_path)
        else:
            return None
        
        if content and not content.startswith("Error"):
            findings = self.search_content(content, file_path.name)
            self.findings[file_path.name] = findings
            return findings
        else:
            print(f"  Error: {content}")
            return None

    def analyze_all_documents(self, documents):
        """Analyze all specified documents"""
        for doc_path in documents:
            full_path = self.base_path / doc_path
            if full_path.exists():
                self.analyze_document(full_path)
            else:
                print(f"File not found: {doc_path}")

    def generate_summary(self):
        """Generate a comprehensive summary of findings"""
        all_aws_services = set()
        all_security_measures = set()
        all_excerpts = []
        
        for file_name, findings in self.findings.items():
            all_aws_services.update(findings['aws_services'])
            all_security_measures.update(findings['cloud_security_measures'])
            for excerpt in findings['relevant_excerpts']:
                all_excerpts.append({
                    'file': file_name,
                    'excerpt': excerpt
                })
        
        summary = {
            'total_documents_analyzed': len(self.findings),
            'aws_services_found': sorted(list(all_aws_services)),
            'cloud_security_measures': sorted(list(all_security_measures)),
            'key_excerpts': all_excerpts[:50],  # Top 50 excerpts
            'document_findings': self.findings
        }
        
        return summary

def main():
    base_path = "/Users/pete/Downloads/FedRAMP Docs"
    
    # List of documents to analyze
    documents = [
        "ClearTriage for Government System Security Plan (SSP).docx",
        "Procedures/(AC) ClearTriage Access Control Procedures.docx",
        "Procedures/(AU) ClearTriage Audit And Accountability Procedures.docx",
        "Procedures/(CM) ClearTriage Configuration Management Procedures.docx",
        "Procedures/(SC) ClearTriage System And Communications Protection Procedures.docx",
        "Procedures/(SI) ClearTriage System And Information Integrity Procedures.docx",
        "Procedures/(IR) ClearTriage Incident Response Procedures.docx",
        "SSP - Appendix N - Continuous Monitoring Plan.xlsx",
        "SSP - Appendix I - Incident Response Plan.docx"
    ]
    
    analyzer = FedRAMPDocumentAnalyzer(base_path)
    analyzer.analyze_all_documents(documents)
    
    summary = analyzer.generate_summary()
    
    # Save findings to JSON
    output_file = "/Users/pete/workspace/fedramp-training/fedramp_aws_findings.json"
    with open(output_file, 'w') as f:
        json.dump(summary, f, indent=2)
    
    # Print summary
    print("\n" + "="*80)
    print("SUMMARY OF AWS AND CLOUD SECURITY FINDINGS")
    print("="*80)
    print(f"\nTotal documents analyzed: {summary['total_documents_analyzed']}")
    print(f"\nAWS Services Found ({len(summary['aws_services_found'])}):")
    for service in summary['aws_services_found']:
        print(f"  - {service}")
    
    print(f"\nCloud Security Measures ({len(summary['cloud_security_measures'])}):")
    for measure in summary['cloud_security_measures'][:20]:  # Show top 20
        print(f"  - {measure}")
    
    print(f"\n\nFull findings saved to: {output_file}")

if __name__ == "__main__":
    main()